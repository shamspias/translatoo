import docx
from googletrans import Translator
import time


def language_translation(source_doc, target_doc, source_lan="en", target_lan="de"):
    """
    To translate Language
    return docx file
    """
    source = source_doc.strip()
    source_language = source_lan.strip()
    target = target_doc.strip()
    target_language = target_lan.strip()

    doc = docx.Document(source + ".docx")
    paragraphs = [para.text for para in doc.paragraphs]
    print(len(paragraphs))
    translator = Translator()
    doc = docx.Document()
    for i, para in enumerate(paragraphs):
        try:
            translation = translator.translate(para, src=source_language, dest=target_language)
            doc.add_paragraph(translation.text)
            time.sleep(1)
            print("Success " + str(i))
        except:
            print("Error " + str(i))
    doc.save(target + ".docx")
    print("Document translation is completed.")
    return True
