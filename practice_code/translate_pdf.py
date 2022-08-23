from pdf2docx import parse
from docx2pdf import convert
import os

import docx
from googletrans import Translator
import time


def language_translation(source_doc, target_doc, source_lan="en", target_lan="de"):
    """
    To translate Language
    return docx file
    """
    source = source_doc.strip()
    source_language = source_lan.strip()
    target = target_doc.strip()
    target_language = target_lan.strip()

    doc = docx.Document(source + "docx")
    paragraphs = [para.text for para in doc.paragraphs]
    print(len(paragraphs))
    translator = Translator()
    doc = docx.Document()
    for i, para in enumerate(paragraphs):
        try:
            translation = translator.translate(para, src=source_language, dest=target_language)
            doc.add_paragraph(translation.text)
            time.sleep(1)
            print("Success " + str(i))
        except:
            print("Error " + str(i))
    doc.save(target + ".docx")
    print("Document translation is completed.")
    return True


def translate_pdf(pdf_file_name, source_ln, target_ln):
    """
    make pdf into word
    translate the word make it into pdf and remove the converted word
    """
    word_file = target_ln + "_" + pdf_file_name[:-3]
    word_file_for_docx = word_file + "docx"
    parse(pdf_file_name, word_file_for_docx, start=0, end=None)
    target_word_file = pdf_file_name[:-3] + "docx"

    language_translation(word_file, target_word_file, source_ln, target_ln)

    new_pdf_file_name = target_ln + "_" + pdf_file_name
    convert(target_word_file, new_pdf_file_name)

    os.remove(word_file)
    os.remove(target_word_file)
    directory = os.getcwd()
    file = directory + new_pdf_file_name
    print(file)
    return file


print(translate_pdf("test.pdf", "en", "bn"))
