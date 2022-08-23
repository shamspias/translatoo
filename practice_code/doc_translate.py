import docx
from googletrans import Translator
import time


def translate_language():
    print("Please enter source document name.")
    source = input().strip()
    print("Please enter source language code.")
    sourceLanguageCode = input().strip()
    print("Please enter target document name.")
    target = input().strip()
    print("Please enter target language code.")
    targetLanguageCode = input().strip()

    doc = docx.Document(source + ".docx")
    paragraphs = [para.text for para in doc.paragraphs]
    print(len(paragraphs))
    translator = Translator()
    doc = docx.Document()
    for i, para in enumerate(paragraphs):
        try:
            translation = translator.translate(para, src=sourceLanguageCode, dest=targetLanguageCode)
            doc.add_paragraph(translation.text)
            time.sleep(1)
            print("Success " + str(i))
        except:
            print("Error " + str(i))
    doc.save(target + ".docx")
    print("Document translation is completed.")
